#!/usr/bin/env python3
"""Fill the TOC field in a DOCX with a page-numbered contents list.

Why this exists: LibreOffice cannot update Word field results headlessly, so a
TOC field alone renders as "Right-click to update…" in the PDF we ship.

So we keep BOTH halves of the field:
  * the w:fldChar/w:instrText field definition stays intact — Word (and the
    house format) still sees a real, updatable TOC that refreshes on F9 or on
    open via w:updateFields; and
  * the field RESULT region (between the 'separate' and 'end' fldChars) is
    filled with the computed entries, which is exactly what a cached field
    result is for — so LibreOffice renders a correct contents list.

Page numbers are read from the already-rendered PDF (same basename). Assumes the
TOC occupies a single page in both passes (true when entry count is small), so
body page numbers do not shift between passes.

Usage: python3 inject_toc.py <docx> <pdf> [maxlevel]
"""
import sys, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from pypdf import PdfReader

def norm(s):
    return re.sub(r'[^a-z0-9]', '', s.lower())

def main():
    docx_path, pdf_path = sys.argv[1], sys.argv[2]
    maxlevel = int(sys.argv[3]) if len(sys.argv) > 3 else 2

    doc = Document(docx_path)

    # 1) collect heading paragraphs (in order) up to maxlevel
    heads = []
    for p in doc.paragraphs:
        sn = p.style.name
        if sn.startswith("Heading"):
            try:
                lvl = int(sn.split()[-1])
            except ValueError:
                continue
            if lvl <= maxlevel and p.text.strip():
                heads.append((lvl, p.text.strip()))

    # 2) page text from the rendered PDF
    reader = PdfReader(pdf_path)
    pages = [norm(pg.extract_text() or "") for pg in reader.pages]

    # 3) map each heading -> first page (>= cursor) whose text contains it
    entries = []
    cursor = 0
    for lvl, text in heads:
        key = norm(text)[:24]
        page = None
        for i in range(cursor, len(pages)):
            if key and key in pages[i]:
                page = i + 1  # 1-based, matches footer "Page N"
                cursor = i
                break
        if page is None:  # fallback: search from start
            for i in range(len(pages)):
                if key and key in pages[i]:
                    page = i + 1
                    break
        entries.append((lvl, text, page or 1))

    # 3b) a multi-page static TOC pushes every body page down between passes
    #     (the pass-1 PDF was measured with the ~1-page placeholder). Offset by
    #     the extra pages the injected TOC occupies. offset is 0 for small TOCs
    #     (e.g. the Lesson Plan), so this never regresses single-page documents.
    LINES_PER_PAGE = 44
    toc_lines = len(entries) + 2            # "Contents" heading + spacing
    toc_pages = max(1, -(-toc_lines // LINES_PER_PAGE))
    offset = toc_pages - 1
    if offset:
        entries = [(lvl, text, pg + offset) for lvl, text, pg in entries]

    # 4) find the TOC field paragraph
    placeholder = None
    for p in doc.paragraphs:
        xml = p._p.xml
        if ("Update Field" in p.text) or ("TOC " in xml and 'instrText' in xml) or ("fldSimple" in xml and "TOC" in xml):
            placeholder = p
            break
    if placeholder is None:
        print("  [inject_toc] no TOC placeholder found in", docx_path); return

    # 5) locate the field's RESULT region: the runs between the 'separate' and
    #    'end' fldChars. Filling that region keeps the field definition intact
    #    (Word can still refresh it) while giving LibreOffice something to draw.
    anchor = placeholder._p
    TYPE = qn('w:fldCharType')
    # prodoc._field() emits begin/instrText/separate/result/end inside ONE run,
    # so search the fldChar ELEMENTS (wherever they sit) rather than the runs.
    sep_el = end_el = None
    for fc in anchor.iter(qn('w:fldChar')):
        kind = fc.get(TYPE)
        if kind == 'separate':
            sep_el = fc
        elif kind == 'end' and sep_el is not None:
            end_el = fc
            break

    GREY = RGBColor(0x33, 0x33, 0x33)

    def style_entry(para, lvl, text, page, text_run=True):
        """Apply TOC-line formatting; text_run=False sets paragraph format only."""
        pf = para.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        if lvl >= 2:
            pf.left_indent = Inches(0.3)
        pf.space_after = Pt(3)
        if not text_run:
            return
        r = para.add_run(text + "\t" + str(page))
        r.font.size = Pt(11 if lvl == 1 else 10.5)
        r.font.name = "Arial"
        r.bold = (lvl == 1)
        r.font.color.rgb = GREY

    if sep_el is not None and end_el is not None:
        from docx.text.paragraph import Paragraph
        holder = sep_el.getparent()          # the run carrying the fldChars

        # Drop the cached placeholder text between 'separate' and 'end'.
        nxt = sep_el.getnext()
        while nxt is not None and nxt is not end_el:
            following = nxt.getnext()
            holder.remove(nxt)
            nxt = following

        # First entry becomes the field's cached RESULT, so the field paragraph
        # itself renders as TOC line 1 in LibreOffice.
        lvl0, text0, page0 = entries[0]
        t = holder.makeelement(qn('w:t'), {})
        t.set(qn('xml:space'), 'preserve')
        t.text = f"{text0}\t{page0}"
        end_el.addprevious(t)

        # Give the field paragraph the same tab/indent treatment as the rest.
        scratch_p = anchor.makeelement(qn('w:p'), {})
        anchor.addprevious(scratch_p)
        style_entry(Paragraph(scratch_p, placeholder._parent), lvl0, "", 0, text_run=False)
        ppr = scratch_p.find(qn('w:pPr'))
        if ppr is not None:
            old = anchor.find(qn('w:pPr'))
            if old is not None:
                anchor.remove(old)
            scratch_p.remove(ppr)
            anchor.insert(0, ppr)
        anchor.getparent().remove(scratch_p)

        # Remaining entries follow as ordinary paragraphs after the field.
        prev = anchor
        for lvl, text, page in entries[1:]:
            new_p = anchor.makeelement(qn('w:p'), {})
            prev.addnext(new_p)
            style_entry(Paragraph(new_p, placeholder._parent), lvl, text, page)
            prev = new_p
        mode = "field result filled (Word can still refresh)"
    else:
        # No usable field region — fall back to plain paragraphs.
        from docx.text.paragraph import Paragraph
        for lvl, text, page in entries:
            new_p = anchor.makeelement(qn('w:p'), {})
            anchor.addprevious(new_p)
            style_entry(Paragraph(new_p, placeholder._parent), lvl, text, page)
        anchor.getparent().remove(anchor)
        mode = "static (no field region found)"

    doc.save(docx_path)
    print(f"  [inject_toc] {docx_path}: wrote {len(entries)} TOC entries "
          f"(pages {entries[0][2]}..{entries[-1][2]}) — {mode}")

if __name__ == "__main__":
    main()
